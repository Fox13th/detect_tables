import os
import pathlib

from docx import Document


def iter_tables(table):
    for row in table.rows:
        for cell in row.cells:
            for nested_table in cell.tables:
                yield nested_table
                yield from iter_tables(nested_table)


def main():
    doc = Document('test.docx')
    # последовательность всех таблиц документа
    all_tables = doc.tables

    list_add_table = []
    for tab in all_tables:
        for n_table in iter_tables(tab):
            list_add_table.append(n_table)
            print("found a nested table %s" % n_table)

    all_tables = all_tables + list_add_table

    print('Всего таблиц в документе:', len(all_tables))

    # создаем пустой словарь под данные таблиц
    data_tables = {i: None for i in range(len(all_tables))}
    # проходимся по таблицам
    for i, table in enumerate(all_tables):
        print('\nДанные таблицы №', i)
        # создаем список строк для таблицы `i` (пока пустые)
        data_tables[i] = [[] for _ in range(len(table.rows))]
        # проходимся по строкам таблицы `i`
        for j, row in enumerate(table.rows):
            # проходимся по ячейкам таблицы `i` и строки `j`
            for cell in row.cells:
                # добавляем значение ячейки в соответствующий
                # список, созданного словаря под данные таблиц
                data_tables[i][j].append(cell.text)

        print(data_tables[i])
        print('\n')

    print('Данные всех таблиц документа:')
    print(data_tables)


def long_name_file(path_name):
    path_name = f'{path_name}'
    return path_name


if __name__ == "__main__":
    path = "C:\\Users\\deanw\\PycharmProjects\\pythonProject\\Users\\попова\\Desktop\\с рабочего стола " \
           "12.12.2019\\Возражения\\ДЖКХ\\Для нач отдела\\из папки " \
           "Махдиева\\УДС\\ЖКХ нормативные акты по нормативам "
    path_name = long_name_file(path)
    pathlib.Path(path_name).mkdir(parents=True, exist_ok=True)

    path_name = os.path.join(path_name, "Пост Адм Ростова нД от 29.01.2016 N 72 Об утв треб к порядку разр и принятия мун правовых актов о нормировании.docx")
    print(path_name)
    #pathlib.Path(path_name).rename('.\\1\\Пост Адм Ростова нД от 29.01.2016 N 72 Об утв треб к порядку разр и принятия мун правовых актов о нормировании.docx')
    os.rename(
        f'{path_name}',
        '.\\1\\Пост Адм Ростова нД от 29.01.2016 N 72 Об утв треб к порядку разр и принятия мун правовых актов о нормировании.docx')
    #main()

#https://habr.com/ru/articles/307186/